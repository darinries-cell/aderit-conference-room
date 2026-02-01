import os
import streamlit as st
from datetime import datetime
import httpx
import json
import re
import concurrent.futures

def get_secret(key):
    try:
        return st.secrets[key]
    except:
        return os.getenv(key)

SUPABASE_URL = "https://yntsgehumjnoxferoycy.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InludHNnZWh1bWpub3hmZXJveWN5Iiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc2OTc0MDA0NywiZXhwIjoyMDg1MzE2MDQ3fQ.OCv9jKm_uDBDFDAgcJlljCAL0XJovK_ihrvA0zUr9qM"

ANTHROPIC_API_KEY = get_secret("ANTHROPIC_API_KEY")
OPENAI_API_KEY = get_secret("OPENAI_API_KEY")
GOOGLE_API_KEY = get_secret("GOOGLE_API_KEY")

import anthropic
import openai
from google import genai

claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)
gemini_client = genai.Client(api_key=GOOGLE_API_KEY)

OLLAMA_AVAILABLE = False
try:
    import ollama
    ollama_client = ollama.Client()
    ollama_client.list()
    OLLAMA_AVAILABLE = True
except:
    pass

def supabase_request(method, table, data=None, params=None):
    url = f"{SUPABASE_URL}/rest/v1/{table}"
    headers = {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation"
    }
    with httpx.Client(timeout=60) as client:
        if method == "GET":
            response = client.get(url, headers=headers, params=params)
        elif method == "POST":
            response = client.post(url, headers=headers, json=data)
        elif method == "PATCH":
            response = client.patch(url, headers=headers, json=data, params=params)
        elif method == "DELETE":
            response = client.delete(url, headers=headers, params=params)
        if response.status_code >= 400:
            st.error(f"DB error: {response.status_code} - {response.text}")
            return None
        if response.text:
            return response.json()
        return []

# ============================================
# FILE / DOCUMENT HANDLING
# ============================================

def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.type
    content = ""
    try:
        if file_type == "text/plain" or uploaded_file.name.endswith('.txt'):
            content = uploaded_file.read().decode('utf-8')
        elif file_type == "text/markdown" or uploaded_file.name.endswith('.md'):
            content = uploaded_file.read().decode('utf-8')
        elif file_type == "text/csv" or uploaded_file.name.endswith('.csv'):
            content = uploaded_file.read().decode('utf-8')
        elif file_type == "application/pdf" or uploaded_file.name.endswith('.pdf'):
            try:
                import pypdf
                reader = pypdf.PdfReader(uploaded_file)
                content = "\n".join([page.extract_text() or "" for page in reader.pages])
            except ImportError:
                content = "[PDF support requires pypdf]"
        elif uploaded_file.name.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(uploaded_file)
                content = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                content = "[DOCX support requires python-docx]"
        elif uploaded_file.name.endswith('.xlsx'):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(uploaded_file, read_only=True)
                sheets = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        rows.append(",".join([str(c) if c is not None else "" for c in row]))
                    sheets.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
                content = "\n\n".join(sheets)
                wb.close()
            except ImportError:
                content = "[XLSX support requires openpyxl]"
        else:
            try:
                content = uploaded_file.read().decode('utf-8')
            except:
                content = f"[Cannot read file type: {file_type}]"
    except Exception as e:
        content = f"[Error reading file: {str(e)[:100]}]"
    return content

def get_doc_save_name(original_filename):
    """Get the right filename for saving to documents. XLSX becomes CSV since we store as text."""
    if original_filename.endswith('.xlsx'):
        return original_filename.rsplit('.', 1)[0] + '.csv'
    return original_filename

def create_docx_buffer(title, content, subtitle=""):
    """Convert text/markdown content into a formatted .docx file. Returns bytes."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import io
    
    doc = DocxDocument()
    
    # Page setup - US Letter
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.style.font.size = Pt(12)
        sub.style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Add generation timestamp
    from datetime import datetime
    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    meta.style.font.size = Pt(9)
    meta.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_paragraph("")  # spacer
    
    # Parse content - handle markdown-style formatting
    lines = content.split('\n')
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            if in_table and table_rows:
                _flush_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue
        
        # Table detection (pipe-delimited)
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            # Skip separator rows like |---|---|
            if all(c in '-| :' for c in stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue
        
        # Flush any pending table
        if in_table and table_rows:
            _flush_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Headings
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            para = doc.add_paragraph(stripped[2:], style='List Bullet')
        # Numbered lists
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):' :
            text = stripped.split(' ', 1)[1] if ' ' in stripped else stripped
            para = doc.add_paragraph(text, style='List Number')
        elif len(stripped) > 3 and stripped[:2].isdigit() and stripped[2] in '.):':
            text = stripped.split(' ', 1)[1] if ' ' in stripped else stripped
            para = doc.add_paragraph(text, style='List Number')
        # Horizontal rule
        elif stripped in ('---', '***', '___'):
            doc.add_paragraph('_' * 50)
        # Bold line (treat as sub-heading)
        elif stripped.startswith('**') and stripped.endswith('**'):
            para = doc.add_paragraph()
            run = para.add_run(stripped.strip('*').strip())
            run.bold = True
            run.font.size = Pt(12)
        # Regular paragraph with inline formatting
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)
    
    # Flush final table
    if table_rows:
        _flush_table(doc, table_rows)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _flush_table(doc, rows):
    """Add a table to the document from parsed rows."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    if not rows:
        return
    
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if i == 0:
                            run.bold = True
    
    doc.add_paragraph("")  # spacer after table

def _add_formatted_text(para, text):
    """Parse inline markdown formatting and add runs to paragraph."""
    import re as _re
    # Split on bold and italic markers
    parts = _re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = para.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Consolas'
            from docx.shared import Pt, RGBColor
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
        else:
            para.add_run(part)

def get_documents(session_id):
    result = supabase_request("GET", "documents", params={
        "select": "*",
        "session_id": f"eq.{session_id}",
        "order": "created_at.asc"
    })
    return result or []

def save_document(session_id, name, content, doc_type="output"):
    result = supabase_request("POST", "documents", data={
        "session_id": session_id,
        "name": name,
        "doc_type": doc_type,
        "content": content
    })
    return result[0] if result else None

def delete_document(doc_id):
    supabase_request("DELETE", "documents", params={"id": f"eq.{doc_id}"})

# ============================================
# PROJECT MANAGEMENT
# ============================================

def get_projects():
    result = supabase_request("GET", "projects", params={"select": "*", "order": "name.asc"})
    return result or []

def create_project(name, emoji="📁", description=""):
    result = supabase_request("POST", "projects", data={"name": name, "emoji": emoji, "description": description})
    if result and len(result) > 0:
        return result[0]["id"]
    return None

def update_project(project_id, updates):
    supabase_request("PATCH", "projects", data=updates, params={"id": f"eq.{project_id}"})

def delete_project(project_id):
    supabase_request("PATCH", "sessions", data={"project_id": None}, params={"project_id": f"eq.{project_id}"})
    supabase_request("DELETE", "projects", params={"id": f"eq.{project_id}"})

# ============================================
# SESSION MANAGEMENT
# ============================================

def get_sessions(project_id=None):
    params = {"select": "*", "order": "updated_at.desc", "limit": "100"}
    if project_id:
        params["project_id"] = f"eq.{project_id}"
    result = supabase_request("GET", "sessions", params=params)
    return result or []

def get_unassigned_sessions():
    result = supabase_request("GET", "sessions", params={"select": "*", "order": "updated_at.desc", "project_id": "is.null", "limit": "100"})
    return result or []

def create_session(name, project_id=None, mode="panel", facilitator_llm="Claude"):
    data = {"name": name[:100], "status": "active", "mode": mode, "facilitator_llm": facilitator_llm}
    if project_id:
        data["project_id"] = project_id
    result = supabase_request("POST", "sessions", data=data)
    if result and len(result) > 0:
        return result[0]["id"]
    return None

def update_session(session_id, updates):
    supabase_request("PATCH", "sessions", data=updates, params={"id": f"eq.{session_id}"})

def delete_session(session_id):
    supabase_request("DELETE", "documents", params={"session_id": f"eq.{session_id}"})
    supabase_request("DELETE", "messages", params={"session_id": f"eq.{session_id}"})
    supabase_request("DELETE", "sessions", params={"id": f"eq.{session_id}"})

def get_messages(session_id):
    result = supabase_request("GET", "messages", params={
        "select": "*",
        "session_id": f"eq.{session_id}",
        "order": "created_at.asc"
    })
    return result or []

def add_message(session_id, role, content, llm_name=None, persona_name=None, persona_emoji=None, round_num=None, message_type="discussion"):
    supabase_request("POST", "messages", data={
        "session_id": session_id,
        "role": role,
        "llm_name": llm_name,
        "persona_name": persona_name,
        "persona_emoji": persona_emoji,
        "content": content,
        "round_num": round_num,
        "message_type": message_type
    })

# ============================================
# ROLES
# ============================================

DEFAULT_ROLES = {
    "positioning_strategist": {"name": "Positioning Strategist", "emoji": "🎯", "category": "Internal - Marketing", "description": "You are the Positioning Strategist. Background: Led positioning for enterprise SaaS including category-creation plays. Studied under April Dunford. Expertise: Why us, why now, category creation, strategic narrative. Style: Strategic, precise, obsessed with differentiation."},
    "copywriter": {"name": "Copywriter", "emoji": "✍️", "category": "Internal - Marketing", "description": "You are the Copywriter. Background: 15 years B2B enterprise copy for Salesforce, Workday, HR tech startups. Expertise: Headlines, website copy, sales decks, one-pagers. Style: Direct, punchy, allergic to jargon."},
    "creative_director": {"name": "Creative Director", "emoji": "🎨", "category": "Internal - Marketing", "description": "You are the Creative Director. Background: Former agency creative lead for enterprise tech brands. Rebranded B2B companies from boring to memorable. Expertise: Brand voice, emotional resonance, differentiation. Style: Creative, provocative, hates forgettable."},
    "buyer_researcher": {"name": "Buyer Researcher", "emoji": "👥", "category": "Internal - Marketing", "description": "You are the Buyer Researcher. Background: Former head of customer research at Gong. Conducted hundreds of buyer interviews. Expertise: Buyer pain language, objections, purchase triggers. Style: Grounded, evidence-based, customer-focused."},
    "ceo": {"name": "CEO Advisor", "emoji": "👔", "category": "Internal - Executive", "description": "You are the CEO Advisor. Background: Founded and sold HR data integration company for $185M. Expertise: Strategy, fundraising, M&A, board management, exit timing. Style: Strategic, sees whole board."},
    "cro": {"name": "CRO Advisor", "emoji": "💼", "category": "Internal - Executive", "description": "You are the CRO Advisor. Background: Former VP Sales at Ceridian, $8M to $80M+ ARR. Expertise: Enterprise sales, hiring, pricing, procurement. Style: Direct, numbers-driven."},
    "chro": {"name": "CHRO (Prospect)", "emoji": "👩‍💼", "category": "Client - Executive", "description": "You are a CHRO at Fortune 500, 15K employees. Situation: 12+ disconnected HR systems, 40% time on data reconciliation. Concerns: Integration complexity, change management, ROI proof, vendor stability. Style: Strategic but practical, been burned before."},
    "cfo_prospect": {"name": "CFO (Prospect)", "emoji": "💰", "category": "Client - Executive", "description": "You are a CFO at $500M company evaluating HR tech. Situation: CHRO wants transformation budget but you've seen IT projects fail. Concerns: TCO, timeline, productivity gains, contract terms. Style: Skeptical, numbers-focused."},
    "talent_acquisition_director": {"name": "Director of Talent Acquisition", "emoji": "🎯", "category": "Client - HR Operations", "description": "You are Director of TA at fast-growing tech company. Situation: ATS and HRIS don't connect, drowning in spreadsheets. Pain: Manual entry, duplicates, slow hiring, poor reporting. Style: Pragmatic, needs solutions NOW."},
    "hr_ops_manager": {"name": "HR Operations Manager", "emoji": "⚙️", "category": "Client - HR Operations", "description": "You are HR Ops Manager at 3K employee company. Situation: You make HR systems work together, fixing data errors daily. Pain: Inconsistencies, manual reporting, compliance anxiety. Style: Detail-oriented, cautious."},
    "it_director": {"name": "IT Director", "emoji": "💻", "category": "Client - IT", "description": "You are IT Director for enterprise apps. Situation: HR bought 15 point solutions without consulting you. Concerns: Security, compliance, APIs, SSO, SLAs. Style: Technical, risk-averse."},
    "solutions_architect": {"name": "Dr. Lena Kowalski", "emoji": "🏗️", "category": "Internal - Technical", "description": "You are Dr. Lena Kowalski, Solutions Architect. Background: 15 years designing enterprise HR integration architectures. Has implemented Workday, SAP SuccessFactors, and Oracle HCM at Fortune 500 companies. Deep understanding of the Aderit four-layer architecture (Connectors → Genome → Knowledge Graph → Agent Swarms). Role: Designs the technical specification for each agent use case — swarm composition, processing logic, Genome dependencies, output formats, implementation phases, and open questions. Voice: Technical but accessible. Writes for a CTO or HRIS Director audience. Specific about architecture, honest about complexity. Never hand-waves. Quality standard: Every spec must answer: 'Could Bennett Reddin read this and start building?' If the answer is no, it's not done."},
    "data_analyst": {"name": "Nolan Park", "emoji": "📊", "category": "Internal - Technical", "description": "You are Nolan Park, Data & Integration Analyst. Background: 12 years mapping enterprise HR data landscapes. Has audited the technology stacks of 50+ large enterprises across all major verticals. Knows which systems contain which data, how accessible it is via API, what real-world extraction challenges exist, and what the data quality traps are. Role: For each use case, maps every required data element to its source system categories, identifies the top 5 most common enterprise systems per data domain, assesses data accessibility difficulty, documents data freshness requirements, and flags data quality risks. Voice: Precise, tabular, systems-aware. Speaks in data domains, APIs, and sync frequencies. Knows the difference between what a system CAN expose and what it typically DOES expose. Quality standard: Every data requirement must name specific systems, not categories. 'HRIS' is not a system. 'Workday HCM, SAP SuccessFactors Employee Central, Oracle HCM Cloud, UKG Pro, ADP Workforce Now' — those are systems."},
    "hr_domain_expert": {"name": "Dr. Amara Osei", "emoji": "📋", "category": "Internal - HR Domain", "description": "You are Dr. Amara Osei, HR Domain Expert & Regulatory Specialist. Background: 20+ years as CHRO and VP HR Operations across healthcare, financial services, manufacturing, and professional services. Deep regulatory knowledge: FLSA, FMLA, ACA, OSHA, HIPAA, SOX, EEOC, EU Pay Transparency Directive, Davis-Bacon, state prevailing wage laws. Has lived every pain point these agents solve. Role: Validates that each use case addresses a real, expensive problem the way practitioners actually experience it. Writes the 'from the buyer shoes' problem statement. Provides regulatory context. Identifies edge cases and implementation gotchas that only a 20-year practitioner would know. Voice: Executive-level HR leader. Speaks in business outcomes, not technology features. Uses the language buyers use internally. Quality standard: Every problem statement must make a real VP of HR say 'that is exactly my problem.' If it sounds like a vendor wrote it, rewrite it."},
    "vertical_analyst": {"name": "Priya Mehta", "emoji": "🔍", "category": "Internal - Analysis", "description": "You are Priya Mehta, Vertical Industry Analyst. Background: Former industry analyst covering HR technology across healthcare, financial services, manufacturing, construction, energy & utilities, transportation & logistics, life sciences, professional services, and hospitality. Knows the regulatory bodies, industry-specific systems, workforce composition nuances, and buying patterns per vertical. Role: For vertical-specific use cases, provides industry context, regulatory drivers, vertical-specific systems, and buyer language. For cross-vertical use cases, identifies which verticals feel the most pain and how the messaging should adapt. Voice: Analyst-grade precision. Cites specific regulations, industry benchmarks, and sector dynamics. Avoids generic 'industry-leading' language. Quality standard: A VP of HR at a hospital should read the healthcare version and feel it was written by someone who understands healthcare. Same for manufacturing, financial services, etc."},
    "product_marketing": {"name": "Tyler Brooks", "emoji": "📣", "category": "Internal - Marketing", "description": "You are Tyler Brooks, Product Marketing & Sales Enablement. Background: 10 years B2B enterprise SaaS marketing at companies selling to Fortune 500. Has launched 20+ product lines. Expert at translating technical capabilities into buyer-facing value propositions, demo narratives, and sales tools. Role: Creates the one-pager, value proposition, ROI framework, demo concept, and discovery questions for each use case. Follows all Aderit language rules strictly. Voice: Clear, benefit-led, urgency-aware. Writes headlines a busy executive reads in 3 seconds and understands. Never leads with features — always leads with the problem or outcome. Quality standard: Every one-pager must pass the 'airplane test' — could Darin hand this to a CHRO sitting next to him on a flight, and would they understand what it does and why they should care in under 60 seconds?"},
    "enterprise_ae_validator": {"name": "Marcus", "emoji": "✅", "category": "Validator", "description": "You are Marcus, Enterprise AE Validator. Background: 15 years selling enterprise software to Fortune 500. Has sat through thousands of procurement reviews. Role: Reality-checks every deliverable with: 'Would this actually land in a Fortune 2000 meeting? Would a VP-level buyer take this seriously? Does it survive the procurement committee?' Red flags to catch: Over-promising capabilities, generic language that could describe any vendor, missing ROI quantification, claims that require proof the company doesn't have yet."},
    "customer_success_validator": {"name": "Aisha", "emoji": "🛡️", "category": "Validator", "description": "You are Aisha, Customer Success & Delivery Validator. Background: 12 years implementing enterprise HR technology. Has managed 50+ enterprise go-lives. Role: Reality-checks every deliverable with: 'Can we actually deliver this? Is the implementation timeline realistic? Are we underestimating the data quality challenges?' Red flags to catch: Specs that assume perfect data, timelines that assume no organizational politics, capabilities that require technology that doesn't exist yet, missing security/compliance considerations."},
    "facilitator": {"name": "Facilitator", "emoji": "🎤", "category": "Internal - Operations", "description": "You are a neutral facilitator. Role: Keep discussion productive, summarize, identify agreement/disagreement, drive conclusions. Style: Neutral, organized, outcome-focused."}
}

@st.cache_data(ttl=5)
def load_roles_from_db():
    result = supabase_request("GET", "roles", params={"select": "*"})
    if not result:
        # Empty DB — seed all defaults
        for key, role in DEFAULT_ROLES.items():
            supabase_request("POST", "roles", data={"key": key, "name": role["name"], "emoji": role["emoji"], "category": role["category"], "description": role["description"]})
        result = supabase_request("GET", "roles", params={"select": "*"})
    else:
        # Sync any new defaults that aren't in DB yet
        existing_keys = {row["key"] for row in result}
        for key, role in DEFAULT_ROLES.items():
            if key not in existing_keys:
                supabase_request("POST", "roles", data={"key": key, "name": role["name"], "emoji": role["emoji"], "category": role["category"], "description": role["description"]})
                result.append({"key": key, "name": role["name"], "emoji": role["emoji"], "category": role["category"], "description": role["description"]})
    roles = {}
    for row in (result or []):
        roles[row["key"]] = {"name": row["name"], "emoji": row["emoji"], "category": row["category"], "description": row["description"]}
    return roles if roles else DEFAULT_ROLES

def save_role_to_db(key, name, emoji, category, description):
    supabase_request("DELETE", "roles", params={"key": f"eq.{key}"})
    supabase_request("POST", "roles", data={"key": key, "name": name, "emoji": emoji, "category": category, "description": description})
    st.cache_data.clear()

def delete_role_from_db(key):
    supabase_request("DELETE", "roles", params={"key": f"eq.{key}"})
    st.cache_data.clear()

@st.cache_data(ttl=5)
def load_room_assignments():
    result = supabase_request("GET", "room_assignments", params={"select": "*"})
    assignments = {}
    for row in (result or []):
        assignments[row["llm_name"]] = row["role_key"]
    if not assignments:
        return {"Claude": "positioning_strategist", "ChatGPT": "copywriter", "Gemini": "creative_director", "Haiku": "buyer_researcher"}
    return assignments

def save_room_assignment(llm_name, role_key):
    supabase_request("DELETE", "room_assignments", params={"llm_name": f"eq.{llm_name}"})
    supabase_request("POST", "room_assignments", data={"llm_name": llm_name, "role_key": role_key, "updated_at": datetime.now().isoformat()})
    st.cache_data.clear()

# ============================================
# LLM FUNCTIONS
# ============================================

COMPANY_CONTEXT = """
COMPANY CONTEXT:
Aderit is an AI-powered HR infrastructure platform seeking $5M Series A. 
Vision: "AWS of workforce data" - unifying disconnected enterprise HR systems.
Founder: Darin Ries
"""

def ask_ollama(prompt, role_description="", memory_context=None):
    if not OLLAMA_AVAILABLE:
        return "[Ollama not available in cloud mode]"
    context = role_description + "\n" + COMPANY_CONTEXT if role_description else COMPANY_CONTEXT
    if memory_context:
        context += "\n\n" + memory_context
    try:
        response = ollama_client.chat(model='qwen2.5:7b', messages=[{'role': 'user', 'content': f"{context}\n\n{prompt}"}])
        return response['message']['content']
    except Exception as e:
        return f"[Ollama error: {str(e)[:100]}]"

def ask_claude(prompt, role_description="", memory_context=None):
    context = role_description + "\n" + COMPANY_CONTEXT if role_description else COMPANY_CONTEXT
    if memory_context:
        context += "\n\n" + memory_context
    message = claude_client.messages.create(model="claude-sonnet-4-20250514", max_tokens=4096, messages=[{"role": "user", "content": f"{context}\n\n{prompt}"}])
    return message.content[0].text

def ask_claude_haiku(prompt, role_description="", memory_context=None):
    context = role_description + "\n" + COMPANY_CONTEXT if role_description else COMPANY_CONTEXT
    if memory_context:
        context += "\n\n" + memory_context
    message = claude_client.messages.create(model="claude-haiku-4-5-20251001", max_tokens=4096, messages=[{"role": "user", "content": f"{context}\n\n{prompt}"}])
    return message.content[0].text

def ask_chatgpt(prompt, role_description="", memory_context=None):
    context = role_description + "\n" + COMPANY_CONTEXT if role_description else COMPANY_CONTEXT
    if memory_context:
        context += "\n\n" + memory_context
    response = openai_client.chat.completions.create(model="gpt-4o", max_tokens=4096, messages=[{"role": "user", "content": f"{context}\n\n{prompt}"}])
    return response.choices[0].message.content

def ask_gemini(prompt, role_description="", memory_context=None):
    context = role_description + "\n" + COMPANY_CONTEXT if role_description else COMPANY_CONTEXT
    if memory_context:
        context += "\n\n" + memory_context
    try:
        response = gemini_client.models.generate_content(model="gemini-2.0-flash", contents=f"{context}\n\n{prompt}")
        return response.text
    except Exception as e:
        return f"[Gemini unavailable: {str(e)[:100]}]"

if OLLAMA_AVAILABLE:
    LLM_FUNCTIONS = {'Claude': ask_claude, 'ChatGPT': ask_chatgpt, 'Gemini': ask_gemini, 'Ollama': ask_ollama}
    LLM_LIST = ['Claude', 'ChatGPT', 'Gemini', 'Ollama']
else:
    LLM_FUNCTIONS = {'Claude': ask_claude, 'ChatGPT': ask_chatgpt, 'Gemini': ask_gemini, 'Haiku': ask_claude_haiku}
    LLM_LIST = ['Claude', 'ChatGPT', 'Gemini', 'Haiku']

ROUND_PROMPTS = {
    1: ("Initial Perspectives", 'Darin says: "{question}"\n\nGive your perspective in 2-3 paragraphs.'),
    2: ("React & Debate", 'Topic: "{question}"\n\nPrevious discussion:\n{previous}\n\nReact: agreements, pushback, questions.'),
    3: ("Specific Proposals", 'Topic: "{question}"\n\nDiscussion so far:\n{previous}\n\nPropose 2-3 specific ideas.'),
    4: ("Critique & Refine", 'Topic: "{question}"\n\nProposals:\n{previous}\n\nWhich ideas have merit? Start converging.'),
    5: ("Final Recommendations", 'Topic: "{question}"\n\nFull discussion:\n{previous}\n\nFinal recommendation.'),
}

# ============================================
# FACILITATOR MODES
# ============================================

def run_panel_discussion(user_input, participants, num_rounds, memory_context, session_id, status_container):
    """Original panel discussion - all LLMs talk in rounds"""
    all_responses = {}
    previous_text = ""
    
    for round_num in range(1, num_rounds + 1):
        round_name, prompt_template = ROUND_PROMPTS[round_num]
        status_container.write(f"**Round {round_num}: {round_name}**")
        
        if round_num == 1:
            prompt = prompt_template.format(question=user_input)
        else:
            prompt = prompt_template.format(question=user_input, previous=previous_text)
        
        responses = {}
        for p in participants:
            status_container.write(f"{p['emoji']} {p['name']}...")
            response = p['func'](prompt, p['description'], memory_context)
            responses[p['llm']] = response
            add_message(session_id, "participant", response, 
                       llm_name=p['llm'], persona_name=p['name'], 
                       persona_emoji=p['emoji'], round_num=round_num)
        
        all_responses[round_num] = responses
        previous_text = "\n\n".join([f"{p['name']}: {responses[p['llm']]}" for p in participants])
    
    return all_responses, previous_text

def run_conversational_mode(user_input, participants, facilitator_llm, memory_context, session_id, status_container, documents):
    """Facilitator orchestrates, calling LLMs one at a time"""
    
    # Build context about available team members
    team_info = "AVAILABLE TEAM MEMBERS:\n"
    for p in participants:
        team_info += f"- {p['name']} ({p['llm']}): {p['description'][:200]}...\n"
    
    # Build document context
    doc_info = ""
    if documents:
        doc_info = "\nAVAILABLE DOCUMENTS:\n"
        for doc in documents:
            doc_info += f"- {doc['name']} ({doc['doc_type']}): {doc['content'][:200]}...\n"
    
    facilitator_func = LLM_FUNCTIONS[facilitator_llm]
    conversation_log = []
    max_turns = 10
    
    # Initial facilitator planning
    plan_prompt = f"""{team_info}{doc_info}

USER REQUEST: {user_input}

You are the facilitator. Analyze this request and plan your approach:
1. Who should you consult first and why?
2. What specific question will you ask them?

Respond in this JSON format:
{{"plan": "your overall plan", "first_call": {{"llm": "LLM name", "question": "specific question to ask"}}}}
"""
    
    status_container.write("**🎯 Facilitator planning...**")
    plan_response = facilitator_func(plan_prompt, "You are a skilled facilitator orchestrating a team discussion.", memory_context)
    add_message(session_id, "facilitator", plan_response, llm_name=facilitator_llm, persona_name="Facilitator", persona_emoji="🎯")
    conversation_log.append(f"Facilitator Plan: {plan_response}")
    
    # Try to parse the plan
    try:
        json_match = re.search(r'\{.*\}', plan_response, re.DOTALL)
        if json_match:
            plan_data = json.loads(json_match.group())
        else:
            plan_data = {"first_call": {"llm": participants[0]['llm'], "question": user_input}}
    except:
        plan_data = {"first_call": {"llm": participants[0]['llm'], "question": user_input}}
    
    turn = 0
    while turn < max_turns:
        turn += 1
        
        # Get the next call from facilitator's plan
        if turn == 1:
            next_call = plan_data.get("first_call", {})
            target_llm = next_call.get("llm", participants[0]['llm'])
            question = next_call.get("question", user_input)
        else:
            # Ask facilitator what to do next
            next_prompt = f"""CONVERSATION SO FAR:
{chr(10).join(conversation_log[-6:])}

Based on the responses so far, decide your next action:
1. Call another team member for input
2. Ask follow-up to same person
3. Conclude and synthesize

Respond in JSON:
{{"action": "call|followup|conclude", "llm": "LLM name if calling", "question": "question to ask", "reason": "why this action"}}
"""
            status_container.write("**🎯 Facilitator deciding next step...**")
            next_response = facilitator_func(next_prompt, "You are a skilled facilitator.", memory_context)
            
            try:
                json_match = re.search(r'\{.*\}', next_response, re.DOTALL)
                if json_match:
                    next_data = json.loads(json_match.group())
                else:
                    next_data = {"action": "conclude"}
            except:
                next_data = {"action": "conclude"}
            
            if next_data.get("action") == "conclude":
                status_container.write("**🎯 Facilitator concluding...**")
                break
            
            target_llm = next_data.get("llm", participants[0]['llm'])
            question = next_data.get("question", "Please provide your thoughts.")
            conversation_log.append(f"Facilitator: {next_data.get('reason', 'Continuing discussion')}")
        
        # Find the participant
        target_participant = next((p for p in participants if p['llm'] == target_llm), participants[0])
        
        status_container.write(f"**{target_participant['emoji']} Asking {target_participant['name']}...**")
        
        # Call the target LLM
        full_question = f"Context: {user_input}\n\nQuestion from facilitator: {question}"
        response = target_participant['func'](full_question, target_participant['description'], memory_context)
        
        add_message(session_id, "participant", response, 
                   llm_name=target_participant['llm'], persona_name=target_participant['name'], 
                   persona_emoji=target_participant['emoji'], round_num=turn)
        
        conversation_log.append(f"{target_participant['name']}: {response}")
        
        # Quality check
        quality_prompt = f"""The {target_participant['name']} responded:
{response[:1000]}

Is this response:
1. HIGH QUALITY - useful, on-topic, actionable
2. NEEDS RETRY - unclear, off-topic, or insufficient

Respond with just "HIGH QUALITY" or "NEEDS RETRY: reason"
"""
        quality_check = facilitator_func(quality_prompt, "You evaluate response quality.", None)
        
        if "NEEDS RETRY" in quality_check and turn < max_turns - 1:
            status_container.write(f"**🔄 Facilitator requesting revision...**")
            retry_prompt = f"Your previous response needs improvement. {quality_check}\n\nOriginal question: {question}\n\nPlease try again with more detail and clarity."
            response = target_participant['func'](retry_prompt, target_participant['description'], memory_context)
            add_message(session_id, "participant", f"[Revised] {response}", 
                       llm_name=target_participant['llm'], persona_name=target_participant['name'], 
                       persona_emoji=target_participant['emoji'], round_num=turn)
            conversation_log[-1] = f"{target_participant['name']} (revised): {response}"
    
    return conversation_log

def run_dispatcher_mode(user_input, participants, facilitator_llm, memory_context, session_id, status_container, documents, file_content=None):
    """Facilitator breaks work into chunks and dispatches in parallel"""
    
    facilitator_func = LLM_FUNCTIONS[facilitator_llm]
    
    # Build context about available team members
    team_info = "AVAILABLE WORKERS:\n"
    for p in participants:
        team_info += f"- {p['name']} ({p['llm']})\n"
    
    # Check if we have CSV/spreadsheet data
    is_spreadsheet = file_content and (',' in file_content[:500] or '\t' in file_content[:500])
    
    if is_spreadsheet:
        # Parse rows for distribution
        lines = file_content.strip().split('\n')
        header = lines[0] if lines else ""
        data_rows = lines[1:] if len(lines) > 1 else []
        
        dispatch_prompt = f"""{team_info}

USER REQUEST: {user_input}

SPREADSHEET DATA:
- Header: {header}
- Total rows: {len(data_rows)}

Divide this work among the available workers. Assign roughly equal chunks.

Respond in JSON format:
{{"tasks": [
    {{"llm": "LLM name", "start_row": 0, "end_row": 25, "instruction": "specific instruction"}},
    ...
]}}
"""
    else:
        dispatch_prompt = f"""{team_info}

USER REQUEST: {user_input}

CONTEXT:
{memory_context[:2000] if memory_context else 'No additional context'}

Break this work into parallel tasks for the available workers. Each worker should get a distinct piece of work.

Respond in JSON format:
{{"tasks": [
    {{"llm": "LLM name", "task": "specific task description"}},
    ...
], "merge_instruction": "how to combine results"}}
"""
    
    status_container.write("**🎯 Facilitator planning work distribution...**")
    dispatch_response = facilitator_func(dispatch_prompt, "You are a skilled work dispatcher.", None)
    add_message(session_id, "facilitator", dispatch_response, llm_name=facilitator_llm, persona_name="Dispatcher", persona_emoji="🎯")
    
    # Parse dispatch plan
    try:
        json_match = re.search(r'\{.*\}', dispatch_response, re.DOTALL)
        if json_match:
            dispatch_plan = json.loads(json_match.group())
        else:
            dispatch_plan = {"tasks": [{"llm": p['llm'], "task": user_input} for p in participants]}
    except:
        dispatch_plan = {"tasks": [{"llm": p['llm'], "task": user_input} for p in participants]}
    
    tasks = dispatch_plan.get("tasks", [])
    merge_instruction = dispatch_plan.get("merge_instruction", "Combine all results into a single coherent output.")
    
    status_container.write(f"**📤 Dispatching {len(tasks)} tasks in parallel...**")
    
    results = []
    
    for i, task in enumerate(tasks):
        target_llm = task.get("llm", participants[0]['llm'])
        target_participant = next((p for p in participants if p['llm'] == target_llm), participants[0])
        
        if is_spreadsheet and "start_row" in task:
            start = task.get("start_row", 0)
            end = task.get("end_row", len(data_rows))
            chunk_rows = data_rows[start:end]
            chunk_data = header + "\n" + "\n".join(chunk_rows)
            
            task_prompt = f"""TASK: {task.get('instruction', user_input)}

DATA (rows {start+1} to {end}):
{chunk_data}

Process this data according to the instruction. Output your results clearly."""
        else:
            task_prompt = f"""TASK: {task.get('task', user_input)}

{memory_context if memory_context else ''}

Complete this task thoroughly."""
        
        status_container.write(f"**{target_participant['emoji']} {target_participant['name']} working...**")
        
        response = target_participant['func'](task_prompt, target_participant['description'], None)
        
        add_message(session_id, "participant", response, 
                   llm_name=target_participant['llm'], persona_name=target_participant['name'], 
                   persona_emoji=target_participant['emoji'], round_num=i+1,
                   message_type="dispatch_result")
        
        results.append({
            "llm": target_llm,
            "name": target_participant['name'],
            "result": response
        })
        
        # Quality check with retry
        quality_prompt = f"""Task result from {target_participant['name']}:
{response[:1000]}

Is this HIGH QUALITY or NEEDS RETRY?"""
        
        quality_check = facilitator_func(quality_prompt, "You evaluate work quality.", None)
        
        if "NEEDS RETRY" in quality_check:
            status_container.write(f"**🔄 Retrying {target_participant['name']}'s task...**")
            response = target_participant['func'](task_prompt + "\n\nPrevious attempt was insufficient. Please be more thorough.", target_participant['description'], None)
            results[-1]["result"] = response
            add_message(session_id, "participant", f"[Revised] {response}", 
                       llm_name=target_participant['llm'], persona_name=target_participant['name'], 
                       persona_emoji=target_participant['emoji'], round_num=i+1)
    
    # Merge results
    status_container.write("**🎯 Facilitator merging results...**")
    
    merge_prompt = f"""ORIGINAL REQUEST: {user_input}

MERGE INSTRUCTION: {merge_instruction}

RESULTS FROM WORKERS:
"""
    for r in results:
        merge_prompt += f"\n--- {r['name']} ---\n{r['result']}\n"
    
    merge_prompt += "\n\nCombine these results into a single coherent output. If the user requested separate files, note that. Otherwise merge into one."
    
    merged_result = facilitator_func(merge_prompt, "You merge work results.", None)
    
    return results, merged_result

# ============================================
# PAGE CONFIG
# ============================================

st.set_page_config(page_title="Aderit Conference Room", page_icon="🏢", layout="wide")

if "current_session_id" not in st.session_state:
    st.session_state.current_session_id = None
if "current_project_id" not in st.session_state:
    st.session_state.current_project_id = None
if "enabled_llms" not in st.session_state:
    st.session_state.enabled_llms = LLM_LIST.copy()
if "num_rounds" not in st.session_state:
    st.session_state.num_rounds = 5
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []  # list of {"name": ..., "content": ...}
if "discussion_mode" not in st.session_state:
    st.session_state.discussion_mode = "panel"
if "facilitator_llm" not in st.session_state:
    st.session_state.facilitator_llm = "Claude"

roles = load_roles_from_db()
room_assignments = load_room_assignments()

# ============================================
# SIDEBAR
# ============================================

with st.sidebar:
    st.title("🏢 Conference Room")
    
    if st.button("➕ New Chat", use_container_width=True, type="primary"):
        st.session_state.current_session_id = None
        st.session_state.file_context = ""
        st.session_state.uploaded_filename = ""
        st.rerun()
    
    st.markdown("---")
    
    # Discussion Mode
    st.markdown("### 🎛️ Mode")
    st.session_state.discussion_mode = st.radio(
        "Discussion Mode",
        ["panel", "conversational", "dispatcher"],
        format_func=lambda x: {"panel": "👥 Panel Discussion", "conversational": "💬 Conversational", "dispatcher": "📤 Dispatcher"}[x],
        index=["panel", "conversational", "dispatcher"].index(st.session_state.discussion_mode),
        label_visibility="collapsed"
    )
    
    if st.session_state.discussion_mode != "panel":
        st.session_state.facilitator_llm = st.selectbox(
            "Facilitator",
            LLM_LIST,
            index=LLM_LIST.index(st.session_state.facilitator_llm) if st.session_state.facilitator_llm in LLM_LIST else 0
        )
    
    if st.session_state.discussion_mode == "panel":
        st.session_state.num_rounds = st.slider("Rounds", 1, 5, st.session_state.num_rounds)
    
    st.markdown("---")
    
    # File upload - supports xlsx
    st.markdown("### 📎 Upload")
    uploaded_files = st.file_uploader("Add files", type=['txt', 'md', 'csv', 'pdf', 'docx', 'xlsx'], key="file_uploader", label_visibility="collapsed", accept_multiple_files=True)
    
    if uploaded_files:
        existing_names = {f["name"] for f in st.session_state.uploaded_files}
        for uf in uploaded_files:
            if uf.name not in existing_names:
                with st.spinner(f"Reading {uf.name}..."):
                    content = extract_text_from_file(uf)
                    st.session_state.uploaded_files.append({"name": uf.name, "content": content})
                st.success(f"✓ {uf.name[:25]}")
    
    if st.session_state.uploaded_files:
        for i, f in enumerate(st.session_state.uploaded_files):
            col_f, col_x = st.columns([4, 1])
            with col_f:
                st.caption(f"📄 {f['name'][:25]}")
            with col_x:
                if st.button("✕", key=f"clear_file_{i}"):
                    st.session_state.uploaded_files.pop(i)
                    st.rerun()
        if st.button("✕ Clear All", key="clear_all_files"):
            st.session_state.uploaded_files = []
            st.rerun()
    
    st.markdown("---")
    
    # Projects
    st.markdown("### 📂 Projects")
    projects = get_projects()
    
    with st.expander("➕ New Project"):
        new_proj_name = st.text_input("Name", key="new_proj_name")
        new_proj_emoji = st.text_input("Emoji", value="📁", key="new_proj_emoji")
        if st.button("Create", key="create_proj"):
            if new_proj_name:
                create_project(new_proj_name, new_proj_emoji)
                st.rerun()
    
    for proj in projects:
        proj_sessions = get_sessions(proj["id"])
        is_current_proj = st.session_state.current_project_id == proj["id"]
        
        with st.expander(f"{proj.get('emoji', '📁')} {proj['name']} ({len(proj_sessions)})", expanded=is_current_proj):
            col1, col2 = st.columns(2)
            with col2:
                if st.button("🗑️", key=f"del_proj_{proj['id']}"):
                    delete_project(proj["id"])
                    st.rerun()
            
            for sess in proj_sessions:
                sess_name = sess["name"][:22] + "..." if len(sess["name"]) > 22 else sess["name"]
                is_current = st.session_state.current_session_id == sess["id"]
                
                if st.button(f"{'▶ ' if is_current else ''}{sess_name}", key=f"sess_{sess['id']}", use_container_width=True):
                    st.session_state.current_session_id = sess["id"]
                    st.session_state.current_project_id = proj["id"]
                    st.rerun()
    
    # Unassigned
    unassigned = get_unassigned_sessions()
    if unassigned:
        with st.expander(f"📋 Unassigned ({len(unassigned)})", expanded=True):
            for sess in unassigned[:15]:
                sess_name = sess["name"][:22] + "..." if len(sess["name"]) > 22 else sess["name"]
                is_current = st.session_state.current_session_id == sess["id"]
                
                if st.button(f"{'▶ ' if is_current else ''}{sess_name}", key=f"sess_{sess['id']}", use_container_width=True):
                    st.session_state.current_session_id = sess["id"]
                    st.session_state.current_project_id = None
                    st.rerun()
    
    st.markdown("---")
    
    # Participants tab
    tab1, tab2, tab3 = st.tabs(["👥", "📚", "⚙️"])
    
    with tab1:
        st.markdown("### Participants")
        role_options = {k: f"{v['emoji']} {v['name']}" for k, v in roles.items()}
        role_keys = list(role_options.keys())
        
        enabled_llms = []
        for llm in LLM_LIST:
            col_check, col_select = st.columns([1, 3])
            with col_check:
                enabled = st.checkbox(f"Enable {llm}", value=llm in st.session_state.enabled_llms, key=f"enable_{llm}", label_visibility="collapsed")
            with col_select:
                current = room_assignments.get(llm, role_keys[0] if role_keys else None)
                current_idx = role_keys.index(current) if current in role_keys else 0
                selected = st.selectbox(f"Role for {llm}", role_keys, index=current_idx, format_func=lambda x: role_options.get(x, x), key=f"assign_{llm}", disabled=not enabled, label_visibility="collapsed")
                if selected != room_assignments.get(llm):
                    save_room_assignment(llm, selected)
                    room_assignments[llm] = selected
            if enabled:
                enabled_llms.append(llm)
        
        st.session_state.enabled_llms = enabled_llms
    
    with tab2:
        st.markdown("### Roles")
        for role_key, role in list(roles.items())[:8]:
            with st.expander(f"{role.get('emoji', '👤')} {role.get('name', role_key)[:15]}"):
                st.caption(role.get('description', '')[:100] + "...")
    
    with tab3:
        st.markdown("### Settings")
        if OLLAMA_AVAILABLE:
            st.success("🖥️ Local")
        else:
            st.info("☁️ Cloud")
        
        if st.button("🔄 Refresh"):
            st.cache_data.clear()
            st.rerun()

# ============================================
# MAIN AREA
# ============================================

col1, col2 = st.columns([5, 1])
with col1:
    st.title("🏢 Aderit Conference Room")
with col2:
    if st.button("🔄 Sync"):
        st.rerun()

# Show mode and room
enabled_llms = st.session_state.enabled_llms
mode_display = {"panel": "👥 Panel", "conversational": "💬 Conversational", "dispatcher": "📤 Dispatcher"}[st.session_state.discussion_mode]

if enabled_llms:
    room_display = " | ".join([f"{roles.get(room_assignments.get(llm, ''), {}).get('emoji', '👤')} {roles.get(room_assignments.get(llm, ''), {}).get('name', '?')[:12]}" for llm in enabled_llms])
    st.caption(f"**Mode:** {mode_display} · **Room:** {room_display}")

if st.session_state.uploaded_files:
    file_names = ", ".join([f["name"][:20] for f in st.session_state.uploaded_files])
    st.info(f"📎 {len(st.session_state.uploaded_files)} file(s): {file_names}")

st.markdown("---")

# Document repository for current session
current_session_id = st.session_state.current_session_id

if current_session_id:
    documents = get_documents(current_session_id)
    
    if documents:
        with st.expander(f"📁 Documents ({len(documents)})", expanded=False):
            for doc in documents:
                col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                with col1:
                    st.write(f"{'📥' if doc['doc_type'] == 'input' else '📤'} {doc['name']}")
                with col2:
                    st.download_button("⬇️", doc['content'], file_name=doc['name'], key=f"dl_{doc['id']}")
                with col3:
                    # Offer docx conversion for text-based documents
                    if doc['name'].endswith(('.md', '.txt', '.csv')) or doc['doc_type'] == 'output':
                        docx_name = doc['name'].rsplit('.', 1)[0] + '.docx' if '.' in doc['name'] else doc['name'] + '.docx'
                        docx_bytes = create_docx_buffer(doc['name'], doc['content'])
                        st.download_button("📝", docx_bytes, file_name=docx_name, key=f"dl_docx_{doc['id']}", help="Download as Word")
                with col4:
                    if st.button("🗑️", key=f"del_doc_{doc['id']}"):
                        delete_document(doc['id'])
                        st.rerun()
    
    # Display messages
    messages = get_messages(current_session_id)
    synthesis_content = ""
    
    for msg in messages:
        if msg["role"] == "user":
            st.markdown(f"**🧑 Darin:** {msg['content']}")
        elif msg["role"] == "facilitator":
            with st.expander(f"🎯 Facilitator ({msg.get('llm_name', '')})", expanded=False):
                st.markdown(msg["content"])
        elif msg["role"] == "participant":
            with st.expander(f"{msg.get('persona_emoji', '👤')} {msg.get('persona_name', 'Unknown')} ({msg.get('llm_name', '?')})", expanded=False):
                st.markdown(msg["content"])
                safe_name = (msg.get('persona_name', 'output') or 'output').replace(' ', '_').replace('/', '-')
                col_p1, col_p2 = st.columns(2)
                with col_p1:
                    st.download_button("⬇️ .md", msg["content"], file_name=f"{safe_name}.md", key=f"dl_part_md_{msg['id']}")
                with col_p2:
                    docx_bytes = create_docx_buffer(msg.get('persona_name', 'Output'), msg["content"])
                    st.download_button("⬇️ .docx", docx_bytes, file_name=f"{safe_name}.docx", key=f"dl_part_docx_{msg['id']}")
        elif msg["role"] == "synthesis":
            st.markdown("### 📋 Synthesis")
            st.markdown(msg["content"])
            synthesis_content = msg["content"]
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.download_button("⬇️ .md", msg["content"], file_name="synthesis.md", key=f"dl_synth_md_{msg['id']}")
            with col_s2:
                docx_bytes = create_docx_buffer("Synthesis", msg["content"])
                st.download_button("⬇️ .docx", docx_bytes, file_name="synthesis.docx", key=f"dl_synth_docx_{msg['id']}")
        elif msg["role"] == "decisions":
            st.markdown("### 🔑 Key Decisions")
            st.markdown(msg["content"])
        elif msg["role"] == "deliverable":
            st.markdown(f"### 📄 {msg.get('persona_name', 'Deliverable')}")
            st.markdown(msg["content"][:500] + "..." if len(msg["content"]) > 500 else msg["content"])
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.download_button("⬇️ .md", msg["content"], file_name=f"{msg.get('persona_name', 'output')}.md", key=f"dl_msg_md_{msg['id']}")
            with col_d2:
                docx_bytes = create_docx_buffer(msg.get('persona_name', 'Deliverable'), msg["content"])
                fname = msg.get('persona_name', 'output').replace('.md', '') + '.docx'
                st.download_button("⬇️ .docx", docx_bytes, file_name=fname, key=f"dl_msg_docx_{msg['id']}", type="primary")
        elif msg["role"] == "merged_output":
            st.markdown("### 📤 Merged Output")
            st.markdown(msg["content"][:2000] + "..." if len(msg["content"]) > 2000 else msg["content"])
            col_m1, col_m2 = st.columns(2)
            with col_m1:
                st.download_button("⬇️ .md", msg["content"], file_name="merged_output.md", key=f"dl_merged_md_{msg['id']}")
            with col_m2:
                docx_bytes = create_docx_buffer("Merged Output", msg["content"])
                st.download_button("⬇️ .docx", docx_bytes, file_name="merged_output.docx", key=f"dl_merged_docx_{msg['id']}", type="primary")
    
    if messages:
        st.markdown("---")
        # Session management
        col1, col2, col3 = st.columns(3)
        with col1:
            new_name = st.text_input("Rename chat", value="", placeholder="New name...", key="rename_sess", label_visibility="collapsed")
            if new_name:
                update_session(current_session_id, {"name": new_name})
                st.rerun()
        with col2:
            move_options = ["(Select project)"] + [p["name"] for p in projects]
            move_to = st.selectbox("Move to project", move_options, key="move_sess", label_visibility="collapsed")
            if move_to != "(Select project)":
                target_proj = next((p for p in projects if p["name"] == move_to), None)
                if target_proj:
                    update_session(current_session_id, {"project_id": target_proj["id"]})
                    st.rerun()
else:
    documents = []
    st.markdown("### Start a new discussion")
    st.markdown(f"**Mode:** {mode_display}")
    if st.session_state.discussion_mode == "panel":
        st.caption("All participants discuss in rounds")
    elif st.session_state.discussion_mode == "conversational":
        st.caption(f"Facilitator ({st.session_state.facilitator_llm}) orchestrates the conversation")
    else:
        st.caption(f"Facilitator ({st.session_state.facilitator_llm}) dispatches work in parallel")

st.markdown("---")

# Input
user_input = st.chat_input("What would you like to discuss?")

if user_input:
    if not enabled_llms:
        st.error("Select at least one participant.")
    else:
        # Create session
        if not current_session_id:
            current_session_id = create_session(
                user_input[:100], 
                st.session_state.current_project_id,
                st.session_state.discussion_mode,
                st.session_state.facilitator_llm
            )
            st.session_state.current_session_id = current_session_id
        
        add_message(current_session_id, "user", user_input)
        
        # Save uploaded files as documents (xlsx saved as csv)
        all_file_content = ""
        for uf in st.session_state.uploaded_files:
            doc_name = get_doc_save_name(uf["name"])
            save_document(current_session_id, doc_name, uf["content"], "input")
            all_file_content += f"## FILE: {uf['name']}\n\n{uf['content'][:8000]}\n\n---\n\n"
        
        memory_context = ""
        if all_file_content:
            memory_context = f"## UPLOADED FILES:\n\n{all_file_content}\n\n"
        
        # Build participants
        participants = []
        for llm in enabled_llms:
            role_key = room_assignments.get(llm, "")
            role = roles.get(role_key, {})
            participants.append({
                'llm': llm,
                'name': role.get('name', 'Unknown'),
                'emoji': role.get('emoji', '👤'),
                'description': role.get('description', ''),
                'func': LLM_FUNCTIONS[llm]
            })
        
        st.markdown(f"**🧑 Darin:** {user_input}")
        
        mode = st.session_state.discussion_mode
        
        with st.status(f"🎯 Running {mode} mode...", expanded=True) as status:
            
            if mode == "panel":
                all_responses, previous_text = run_panel_discussion(
                    user_input, participants, st.session_state.num_rounds, 
                    memory_context, current_session_id, st
                )
                final_responses = all_responses[st.session_state.num_rounds]
                
            elif mode == "conversational":
                conversation_log = run_conversational_mode(
                    user_input, participants, st.session_state.facilitator_llm,
                    memory_context, current_session_id, st, documents
                )
                previous_text = "\n\n".join(conversation_log)
                final_responses = {p['llm']: "" for p in participants}
                
            else:  # dispatcher
                results, merged_result = run_dispatcher_mode(
                    user_input, participants, st.session_state.facilitator_llm,
                    memory_context, current_session_id, st, documents, all_file_content if all_file_content else None
                )
                previous_text = merged_result
                final_responses = {r['llm']: r['result'] for r in results}
                
                # Save merged result as document AND as a visible message
                save_document(current_session_id, "merged_output.md", merged_result, "output")
                add_message(current_session_id, "merged_output", merged_result, 
                           llm_name=st.session_state.facilitator_llm, persona_name="Merged Output", persona_emoji="📤")
            
            # Synthesis
            st.write("**Synthesizing...**")
            synth_prompt = f'Topic: "{user_input}"\n\nDiscussion:\n{previous_text[:6000]}\n\nSynthesize in 3-4 paragraphs.'
            synthesis = ask_claude(synth_prompt, "You are a neutral facilitator.", memory_context)
            add_message(current_session_id, "synthesis", synthesis)
            
            decisions_prompt = f"List 3-5 key decisions as bullet points:\n\n{synthesis}"
            key_decisions = ask_claude(decisions_prompt, "Extract key decisions.", None)
            add_message(current_session_id, "decisions", key_decisions)
            
            # Build full report markdown
            report_md = f"# {user_input}\n\n"
            report_md += f"**Date:** {datetime.now().strftime('%B %d, %Y')}\n"
            report_md += f"**Mode:** {mode}\n"
            report_md += f"**Participants:** {', '.join([p['name'] for p in participants])}\n\n"
            report_md += "---\n\n"
            report_md += f"## Synthesis\n\n{synthesis}\n\n"
            report_md += f"## Key Decisions\n\n{key_decisions}\n\n"
            
            # Save individual participant outputs as .md files
            if mode == "panel":
                report_md += "---\n\n## Participant Responses\n\n"
                last_round = st.session_state.num_rounds
                for p in participants:
                    resp = all_responses.get(last_round, {}).get(p['llm'], '')
                    if resp:
                        participant_md = f"# {p['name']} — {p['emoji']}\n\n"
                        participant_md += f"**Role:** {p['description'][:200]}\n\n"
                        participant_md += f"**Topic:** {user_input}\n\n---\n\n{resp}"
                        safe_name = p['name'].replace(' ', '_').replace('/', '-')
                        save_document(current_session_id, f"{safe_name}.md", participant_md, "output")
                        report_md += f"### {p['emoji']} {p['name']}\n\n{resp}\n\n"
            elif mode == "dispatcher":
                report_md += "---\n\n## Merged Output\n\n" + merged_result + "\n\n"
                report_md += "## Individual Worker Results\n\n"
                for r in results:
                    worker_md = f"# {r['name']} — Worker Output\n\n"
                    worker_md += f"**Topic:** {user_input}\n\n---\n\n{r['result']}"
                    safe_name = r['name'].replace(' ', '_').replace('/', '-')
                    save_document(current_session_id, f"{safe_name}.md", worker_md, "output")
                    report_md += f"### {r['name']}\n\n{r['result']}\n\n"
            else:  # conversational
                report_md += "---\n\n## Conversation Log\n\n"
                for entry in conversation_log:
                    report_md += f"{entry}\n\n"
            
            # Save full report as .md
            save_document(current_session_id, "report.md", report_md, "output")
            
            # Check for file request
            if any(kw in user_input.lower() for kw in ["create a file", "create a markdown", "generate a file", "write a file", "create a prompt"]):
                st.write("**Generating deliverable...**")
                file_prompt = f"Based on this discussion:\n{synthesis}\n\nCreate the requested deliverable. Output ONLY the file content."
                deliverable = ask_claude(file_prompt, "Create professional deliverables.", memory_context)
                
                filename = "deliverable.md"
                if "lovable" in user_input.lower():
                    filename = "lovable_prompt.md"
                elif "website" in user_input.lower():
                    filename = "website_spec.md"
                
                add_message(current_session_id, "deliverable", deliverable, persona_name=filename)
                save_document(current_session_id, filename, deliverable, "output")
            
            update_session(current_session_id, {"status": "complete", "updated_at": datetime.now().isoformat()})
            status.update(label="✅ Complete!", state="complete")
        
        # Show dispatcher merged output prominently
        if mode == "dispatcher":
            st.markdown("### 📤 Merged Output")
            st.markdown(merged_result[:2000] + "..." if len(merged_result) > 2000 else merged_result)
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                st.download_button("⬇️ Download .md", merged_result, file_name="merged_output.md", key="dl_merged_md_new", type="primary")
            with col_dl2:
                docx_bytes = create_docx_buffer("Merged Output", merged_result, subtitle=user_input[:100])
                st.download_button("⬇️ Download .docx", docx_bytes, file_name="merged_output.docx", key="dl_merged_docx_new", type="secondary")
        
        st.markdown("### 📋 Synthesis")
        st.markdown(synthesis)
        st.markdown("### 🔑 Key Decisions")
        st.markdown(key_decisions)
        
        # Full report downloads — .md and .docx
        st.markdown("---")
        st.markdown("### 📄 Full Report")
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            st.download_button("⬇️ Report (.md)", report_md, file_name="report.md", key="dl_report_md", type="primary")
        with col_r2:
            docx_report = create_docx_buffer(user_input[:80], report_md, subtitle="Aderit Conference Room Report")
            st.download_button("⬇️ Report (.docx)", docx_report, file_name="report.docx", key="dl_report_docx", type="primary")
        
        st.rerun()
